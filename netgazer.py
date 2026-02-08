#!/usr/bin/env python3
"""
NetGazer - Scan for web servers, capture their screens and place them in an HTML or word document table (docx)
Author: Vahe Demirkhanyan
"""

from __future__ import annotations

import atexit
import base64
import html
import ipaddress
import os
import re
import socket
import ssl
import sys
import tempfile
import time
import logging
import shutil
import errno

from collections import Counter
from concurrent.futures import ThreadPoolExecutor, wait, FIRST_COMPLETED
from datetime import datetime
from io import BytesIO
from itertools import product
from pathlib import Path
from threading import Lock, local
from typing import List, Tuple, Optional, Dict, Iterator


# ── Optional deps (loaded safely so we can print a clean error) ────────
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL
except ImportError:
    docx = None
    Inches = Pt = RGBColor = WD_PARAGRAPH_ALIGNMENT = WD_ALIGN_VERTICAL = None  # type: ignore

try:
    from selenium import webdriver
    from selenium.webdriver.firefox.service import Service as FirefoxService
    from selenium.webdriver.firefox.options import Options
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.common.exceptions import (
        TimeoutException,
        WebDriverException,
        UnexpectedAlertPresentException,
        NoSuchWindowException,
    )
except ImportError:
    webdriver = None
    FirefoxService = Options = WebDriverWait = None  # type: ignore
    TimeoutException = WebDriverException = UnexpectedAlertPresentException = NoSuchWindowException = None  # type: ignore

try:
    from webdriver_manager.firefox import GeckoDriverManager
except ImportError:
    GeckoDriverManager = None  # type: ignore

try:
    from tqdm import tqdm
except ImportError:
    tqdm = None  # type: ignore

try:
    from urllib.parse import urlparse
except ImportError:
    from urlparse import urlparse  # type: ignore

import argparse


# ═══════════════════════════════════════════════════════════════
#  Constants
# ═══════════════════════════════════════════════════════════════

DEFAULT_PORTS: List[int] = [80, 443]
DEFAULT_TIMEOUT: int = 30

MAX_PORT_WORKERS: int = 50
PORT_CHECK_TIMEOUT: int = 3
TLS_PROBE_TIMEOUT: float = 2.0

_TITLE_MAX_LEN = 80
_BROWSER_RESET_EVERY = 25

# Parallel capture: keep modest so we don't swamp RAM with Firefox instances.
# Each headless Firefox ≈ 200-300 MB.  3 is a good default.
_MAX_SHOT_WORKERS = 3

_MAX_FULLPAGE_HEIGHT = 22000


# ═══════════════════════════════════════════════════════════════
#  Interesting-page tag detection
# ═══════════════════════════════════════════════════════════════

_TAG_BADGE_COLORS: Dict[str, str] = {
    "LOGIN":   "#e67e22",
    "ADMIN":   "#e74c3c",
    "API":     "#8e44ad",
    "DEFAULT": "#3498db",
    "403":     "#f39c12",
    "401":     "#f39c12",
    "404":     "#95a5a6",
    "ERROR":   "#c0392b",
}


def _detect_tags(title: str, url: str = "", final_url: str = "") -> List[str]:
    t = (title or "").lower()
    u = (url or "").lower()
    fu = (final_url or "").lower()
    blob = " ".join([t, u, fu])

    tags: List[str] = []

    if any(k in blob for k in ("login", "sign in", "log in", "signin",
                              "authenticate", "password", "credentials",
                              "/login", "/signin", "/auth")):
        tags.append("LOGIN")

    if any(k in blob for k in ("admin", "dashboard", "control panel",
                              "management", "console", "webmin",
                              "phpmyadmin", "grafana", "kibana",
                              "jenkins", "portainer",
                              "/admin", "/wp-admin", "/manager/html")):
        tags.append("ADMIN")

    if any(k in blob for k in ("/api", "swagger", "openapi", "graphql", "/v1/", "/v2/")):
        tags.append("API")

    if any(k in t for k in ("default page", "welcome to nginx",
                            "apache2", "it works", "iis windows",
                            "test page", "default web site",
                            "congratulations", "coming soon",
                            "under construction", "placeholder")):
        tags.append("DEFAULT")

    if "forbidden" in t or "403" in t:
        tags.append("403")
    if "unauthorized" in t or "401" in t:
        tags.append("401")
    if any(k in t for k in ("not found", "404")):
        tags.append("404")
    if any(k in t for k in ("internal server error", "500", "502",
                            "503", "bad gateway", "service unavailable")):
        tags.append("ERROR")

    return tags


# ═══════════════════════════════════════════════════════════════
#  Console helpers
# ═══════════════════════════════════════════════════════════════

def _c_green(t: str) -> str:  return f"\033[92m{t}\033[0m"
def _c_yellow(t: str) -> str: return f"\033[93m{t}\033[0m"
def _c_red(t: str) -> str:    return f"\033[91m{t}\033[0m"
def _c_dim(t: str) -> str:    return f"\033[2m{t}\033[0m"
def _c_bold(t: str) -> str:   return f"\033[1m{t}\033[0m"


def _truncate(s: str, n: int = _TITLE_MAX_LEN) -> str:
    s = s or ""
    return s if len(s) <= n else s[:n - 1] + "…"


def _fmt_elapsed(secs: float) -> str:
    if secs < 60:
        return f"{secs:.1f}s"
    m, s = divmod(secs, 60)
    return f"{int(m)}m {s:.0f}s"


def _dedupe_preserve_order(items: List[str]) -> List[str]:
    seen: set[str] = set()
    out: List[str] = []
    for x in items:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out


def _atomic_replace(src: Path, dst: Path) -> None:
    try:
        src.replace(dst)
    except Exception:
        shutil.move(str(src), str(dst))


def _twrite(msg: str) -> None:
    """Write a line through tqdm (if available) or plain print."""
    try:
        tqdm.write(msg)  # type: ignore[union-attr]
    except Exception:
        try:
            print(msg, flush=True)
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════
#  Logging
# ═══════════════════════════════════════════════════════════════

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════
#  DNS cache
# ═══════════════════════════════════════════════════════════════

_DNS_CACHE: dict[str, list[tuple[int, tuple]]] = {}
_DNS_LOCK = Lock()
_DNS_FAILS: set[str] = set()
_DNS_FAILS_LOCK = Lock()


def _sockaddr_with_port(sockaddr: tuple, port: int) -> tuple:
    if len(sockaddr) == 2:
        return (sockaddr[0], port)
    if len(sockaddr) == 4:
        return (sockaddr[0], port, sockaddr[2], sockaddr[3])
    return sockaddr


def _is_ip_literal(host: str) -> bool:
    try:
        ipaddress.ip_address(host)
        return True
    except ValueError:
        return False


def resolve_host_cached(host: str, prefer_ipv4: bool = True) -> list[tuple[int, tuple]]:
    with _DNS_FAILS_LOCK:
        if host in _DNS_FAILS:
            return []

    with _DNS_LOCK:
        if host in _DNS_CACHE:
            addrs = _DNS_CACHE[host]
            return addrs if not prefer_ipv4 else sorted(addrs, key=lambda x: 0 if x[0] == socket.AF_INET else 1)

    addrs: list[tuple[int, tuple]] = []

    try:
        ip = ipaddress.ip_address(host)
        if ip.version == 4:
            addrs = [(socket.AF_INET, (host, 0))]
        else:
            addrs = [(socket.AF_INET6, (host, 0, 0, 0))]
    except ValueError:
        try:
            infos = socket.getaddrinfo(host, None, type=socket.SOCK_STREAM)
            seen = set()
            for family, _, _, _, sockaddr in infos:
                if family not in (socket.AF_INET, socket.AF_INET6):
                    continue
                key = (family, sockaddr)
                if key not in seen:
                    seen.add(key)
                    addrs.append((family, sockaddr))
        except (socket.gaierror, OSError):
            addrs = []
            with _DNS_FAILS_LOCK:
                _DNS_FAILS.add(host)

    with _DNS_LOCK:
        _DNS_CACHE[host] = addrs

    return addrs if not prefer_ipv4 else sorted(addrs, key=lambda x: 0 if x[0] == socket.AF_INET else 1)


def is_ipv6(addr: str) -> bool:
    try:
        ipaddress.IPv6Address(addr)
        return True
    except ValueError:
        return False


def split_host_port_maybe(s: str) -> Tuple[str, Optional[int]]:
    s = s.strip()
    if s.startswith('[') and ']' in s:
        end = s.index(']')
        host = s[1:end]
        rest = s[end + 1:]
        if rest.startswith(':'):
            port_part = rest[1:]
            if port_part.isdigit():
                return host, int(port_part)
        return host, None

    if s.count(':') == 1:
        left, right = s.rsplit(':', 1)
        if right.isdigit():
            return left, int(right)

    return s, None


# ═══════════════════════════════════════════════════════════════
#  Network / host parsing
# ═══════════════════════════════════════════════════════════════

_IP_RANGEISH_RE = re.compile(r'^[0-9.\-]+$')
_CIDRISH_RE = re.compile(r'^[0-9.:/]+$')


def _looks_like_ip_range(token: str) -> bool:
    return bool(_IP_RANGEISH_RE.match(token)) and "." in token


def _looks_like_cidr(token: str) -> bool:
    return bool(_CIDRISH_RE.match(token)) and "/" in token


class NetworkScanner:
    @staticmethod
    def parse_hosts(input_value: str) -> List[str]:
        input_path = Path(input_value)
        return (NetworkScanner._parse_hosts_file(input_path)
                if input_path.is_file()
                else NetworkScanner._parse_single_input(input_value))

    @staticmethod
    def _parse_hosts_file(filepath: Path) -> List[str]:
        hosts: List[str] = []
        try:
            with filepath.open('r', encoding='utf-8') as f:
                for line_num, line in enumerate(f, 1):
                    line = line.strip()
                    if not line or line.startswith('#'):
                        continue
                    if '#' in line:
                        line = line.split('#', 1)[0].strip()
                        if not line:
                            continue
                    try:
                        hosts.extend(NetworkScanner._parse_single_input(line))
                    except Exception as e:
                        logger.warning(f"Skipping invalid line {line_num} in {filepath}: {e}")
        except IOError as e:
            logger.error(f"Error reading file {filepath}: {e}")
            sys.exit(1)
        return hosts

    @staticmethod
    def _parse_single_input(text: str) -> List[str]:
        raw = text.strip()
        if not raw:
            return []

        host_candidate = raw
        had_scheme = False
        if raw.startswith("//"):
            raw = "http:" + raw
        if '://' in raw:
            had_scheme = True
            u = urlparse(raw)
            host_candidate = u.netloc or u.path

        # Only strip URL path components when input had a scheme.
        # Raw CIDRs like "192.168.1.0/24" must NOT be split here.
        if had_scheme:
            host_candidate = host_candidate.split('/')[0]
        host_candidate = host_candidate.strip().rstrip('.')

        if "@" in host_candidate:
            host_candidate = host_candidate.split("@", 1)[-1]

        if host_candidate.startswith('[') and ']' in host_candidate:
            host_inner = host_candidate[1:host_candidate.index(']')]
            rest = host_candidate[host_candidate.index(']') + 1:]
            if rest.startswith(':'):
                host_candidate = f"{host_inner}{rest}"
            else:
                host_candidate = host_inner

        try:
            ipaddress.ip_address(host_candidate)
            return [host_candidate]
        except ValueError:
            pass

        if '/' in host_candidate and _looks_like_cidr(host_candidate):
            return NetworkScanner._expand_cidr(host_candidate)

        if '-' in host_candidate and _looks_like_ip_range(host_candidate):
            return NetworkScanner._expand_range(host_candidate)

        # Fallback: if no scheme was given and a '/' remains (e.g. "host.com/path"),
        # strip the path part and treat as plain hostname.
        if '/' in host_candidate:
            host_candidate = host_candidate.split('/')[0]

        return [host_candidate]

    @staticmethod
    def _expand_cidr(cidr_str: str) -> List[str]:
        try:
            network = ipaddress.ip_network(cidr_str, strict=False)
            if network.num_addresses == 1:
                return [str(network.network_address)]
            return [str(ip) for ip in network.hosts()]
        except ValueError as e:
            logger.error(f"Invalid CIDR notation '{cidr_str}': {e}")
            return []

    @staticmethod
    def _expand_range(range_str: str) -> List[str]:
        dash_count = range_str.count('-')
        if dash_count == 1:
            return NetworkScanner._expand_simple_range(range_str)
        if dash_count > 1:
            return NetworkScanner._expand_complex_range(range_str)
        return []

    @staticmethod
    def _expand_simple_range(range_str: str) -> List[str]:
        try:
            start_ip_str, end_str = range_str.split('-', 1)
            if '.' not in end_str:
                base = '.'.join(start_ip_str.split('.')[:-1])
                end_ip_str = f"{base}.{end_str}"
            else:
                end_ip_str = end_str

            start_ip = ipaddress.ip_address(start_ip_str)
            end_ip = ipaddress.ip_address(end_ip_str)

            return [str(ipaddress.ip_address(ip_int))
                    for ip_int in range(int(start_ip), int(end_ip) + 1)]
        except (ValueError, ipaddress.AddressValueError) as e:
            logger.error(f"Invalid IP range '{range_str}': {e}")
            return []

    @staticmethod
    def _expand_complex_range(range_str: str) -> List[str]:
        try:
            octet_parts = range_str.split('.')
            if len(octet_parts) != 4:
                raise ValueError("Invalid IP format")

            octet_ranges = []
            for part in octet_parts:
                if '-' in part:
                    start_s, end_s = part.split('-', 1)
                    start_val, end_val = int(start_s), int(end_s)
                    if not (0 <= start_val <= 255 and 0 <= end_val <= 255):
                        raise ValueError("Octet out of range")
                    octet_ranges.append(range(start_val, end_val + 1))
                else:
                    val = int(part)
                    if not (0 <= val <= 255):
                        raise ValueError("Octet out of range")
                    octet_ranges.append([val])

            return ['.'.join(map(str, combo)) for combo in product(*octet_ranges)]
        except (ValueError, IndexError) as e:
            logger.error(f"Invalid complex range '{range_str}': {e}")
            return []


# ═══════════════════════════════════════════════════════════════
#  Port scanner
# ═══════════════════════════════════════════════════════════════

_TLS_CTX_LOCK = Lock()
_TLS_CTX: Optional[ssl.SSLContext] = None


def _get_tls_ctx() -> ssl.SSLContext:
    global _TLS_CTX
    with _TLS_CTX_LOCK:
        if _TLS_CTX is not None:
            return _TLS_CTX
        ctx = ssl.SSLContext(ssl.PROTOCOL_TLS_CLIENT)
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        _TLS_CTX = ctx
        return ctx


class PortScanner:
    @staticmethod
    def probe_port_with_addrs(host: str, port: int, addr_list: List[Tuple[int, tuple]], timeout: int = PORT_CHECK_TIMEOUT) -> Tuple[bool, str]:
        if not addr_list:
            return False, "http"

        sni_name = None if _is_ip_literal(host) else host
        tls_ctx = _get_tls_ctx()

        for family, sockaddr0 in addr_list:
            sockaddr = _sockaddr_with_port(sockaddr0, port)
            sock: Optional[socket.socket] = None
            try:
                sock = socket.socket(family, socket.SOCK_STREAM)
                sock.settimeout(float(timeout))
                rc = sock.connect_ex(sockaddr)
                if rc != 0:
                    try:
                        sock.close()
                    except Exception:
                        pass
                    continue

                scheme = "http"
                try:
                    sock.settimeout(float(min(timeout, TLS_PROBE_TIMEOUT)))
                    with tls_ctx.wrap_socket(sock, server_hostname=sni_name, do_handshake_on_connect=False) as ssock:
                        ssock.settimeout(float(TLS_PROBE_TIMEOUT))
                        ssock.do_handshake()
                    scheme = "https"
                except Exception:
                    try:
                        sock.close()
                    except Exception:
                        pass

                return True, scheme

            except OSError as e:
                try:
                    if sock:
                        sock.close()
                except Exception:
                    pass
                if getattr(e, "errno", None) in (errno.EMFILE, errno.ENOBUFS, errno.EADDRNOTAVAIL, errno.EADDRINUSE):
                    time.sleep(0.01)
                continue

        return False, "http"

    @staticmethod
    def probe_port(host: str, port: int, timeout: int = PORT_CHECK_TIMEOUT) -> Tuple[bool, str]:
        addr_list = resolve_host_cached(host, prefer_ipv4=True)
        return PortScanner.probe_port_with_addrs(host, port, addr_list, timeout)

    @staticmethod
    def scan_host(host: str, ports: List[int]) -> Optional[Tuple[str, List[Tuple[int, str]]]]:
        h, explicit_port = split_host_port_maybe(host)

        addr_list = resolve_host_cached(h, prefer_ipv4=True)
        if not addr_list:
            return None

        if explicit_port is not None:
            ok, scheme = PortScanner.probe_port_with_addrs(h, explicit_port, addr_list)
            return (h, [(explicit_port, scheme)]) if ok else None

        open_ports: List[Tuple[int, str]] = []
        for port in ports:
            ok, scheme = PortScanner.probe_port_with_addrs(h, port, addr_list)
            if ok:
                open_ports.append((port, scheme))

        open_ports.sort(key=lambda x: x[0])
        return (h, open_ports) if open_ports else None


# ═══════════════════════════════════════════════════════════════
#  WebDriver manager
# ═══════════════════════════════════════════════════════════════

def _resolve_geckodriver_path() -> Optional[str]:
    env = (os.environ.get("GECKODRIVER") or os.environ.get("GECKODRIVER_PATH") or "").strip()
    if env:
        p = Path(env)
        if p.is_file():
            return str(p)
    which = shutil.which("geckodriver")
    if which:
        return which
    return None


class WebDriverManager:
    _driver_options_cache: Optional[Options] = None
    _gecko_path_cache: Optional[str] = None
    _lock = Lock()

    _tl = local()
    _drivers: set = set()
    _drivers_lock = Lock()
    _atexit_registered = False

    @classmethod
    def _get_cached_options(cls) -> Options:
        if cls._driver_options_cache is None:
            options = Options()
            options.add_argument("--headless")

            prefs = {
                "dom.webdriver.enabled": False,
                "media.volume_scale": "0.0",
                "security.enterprise_roots.enabled": True,
                "dom.security.https_only_mode": False,
                "dom.security.https_only_mode_pbm": False,
                "dom.webnotifications.enabled": False,
                "dom.push.enabled": False,
                "permissions.default.desktop-notification": 2,
                "geo.enabled": False,
                "browser.privatebrowsing.autostart": True,
                "browser.cache.disk.enable": False,
                "browser.cache.memory.enable": False,
                "network.http.use-cache": False,
                "dom.disable_open_during_load": True,
                "dom.popup_maximum": 0,
                "privacy.trackingprotection.enabled": True,
                "browser.download.manager.showWhenStarting": False,
                "browser.download.alwaysOpenPanel": False,
                "browser.download.manager.useWindow": False,
                "browser.helperApps.alwaysAsk.force": False,
                "webdriver_accept_untrusted_certs": True,
                "accept_untrusted_certs": True,
                "assume_untrusted_cert_issuer": False,
            }
            for pref, value in prefs.items():
                try:
                    options.set_preference(pref, value)
                except Exception:
                    pass

            try:
                options.set_capability("acceptInsecureCerts", True)
            except Exception:
                pass

            cls._driver_options_cache = options
        return cls._driver_options_cache

    @classmethod
    def _get_gecko_path(cls) -> str:
        if cls._gecko_path_cache:
            return cls._gecko_path_cache
        with cls._lock:
            if cls._gecko_path_cache:
                return cls._gecko_path_cache
            gecko_path = _resolve_geckodriver_path()
            if gecko_path is None:
                if GeckoDriverManager is None:
                    raise RuntimeError("geckodriver not found in PATH and webdriver-manager is not installed")
                gecko_path = GeckoDriverManager().install()
            cls._gecko_path_cache = gecko_path
            return gecko_path

    @classmethod
    def warmup(cls) -> None:
        """Pre-cache geckodriver path and options on the main thread
        so worker threads don't all race to resolve them."""
        cls._get_cached_options()
        cls._get_gecko_path()

    @classmethod
    def create_driver(cls, timeout: int = DEFAULT_TIMEOUT) -> webdriver.Firefox:
        d = getattr(cls._tl, "d", None)
        if d is not None:
            try:
                _ = d.current_url
                d.set_page_load_timeout(timeout)
                try:
                    d.set_script_timeout(timeout)
                except Exception:
                    pass
                return d
            except Exception:
                try:
                    d.quit()
                except Exception:
                    pass
                setattr(cls._tl, "d", None)

        try:
            options = cls._get_cached_options()
            gecko_path = cls._get_gecko_path()

            service = FirefoxService(executable_path=gecko_path)
            d = webdriver.Firefox(service=service, options=options)
            d.set_page_load_timeout(timeout)
            try:
                d.set_script_timeout(timeout)
            except Exception:
                pass

            try:
                d.set_window_size(1920, 1080)
            except Exception:
                pass

            try:
                d.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            except Exception:
                pass

            setattr(d, "_ng_nav", 0)
            setattr(cls._tl, "d", d)
            with cls._drivers_lock:
                cls._drivers.add(d)

            if not cls._atexit_registered:
                cls._atexit_registered = True
                atexit.register(cls.cleanup_all)

            return d

        except Exception as e:
            logger.error(f"Failed to create WebDriver: {e}")
            sys.exit(1)

    @classmethod
    def cleanup_all(cls) -> None:
        """Quit every tracked browser instance.  Thread-safe."""
        with cls._drivers_lock:
            ds = list(cls._drivers)
            cls._drivers.clear()
        # Quit in parallel — each quit() can take a second or two
        if len(ds) <= 1:
            for d in ds:
                try:
                    d.quit()
                except Exception:
                    pass
        else:
            def _q(d):
                try:
                    d.quit()
                except Exception:
                    pass
            with ThreadPoolExecutor(max_workers=len(ds)) as ex:
                list(ex.map(_q, ds))

    @staticmethod
    def _dom_has_meaningful_content(driver: webdriver.Firefox) -> bool:
        try:
            return bool(driver.execute_script("""
                const b = document.body;
                if (!b) return false;
                const t = (b.innerText || "").trim();
                if (t.length > 0) return true;
                if (document.images && document.images.length > 0) return true;
                if (document.querySelector("input,button,form,a,canvas,svg,video,iframe")) return true;
                return (b.scrollHeight || 0) > 120;
            """))
        except Exception:
            return False

    @classmethod
    def _maybe_reset_state(cls, driver: webdriver.Firefox) -> None:
        try:
            n = int(getattr(driver, "_ng_nav", 0)) + 1
            setattr(driver, "_ng_nav", n)
        except Exception:
            n = 0
        if not n or n % _BROWSER_RESET_EVERY != 0:
            return
        try:
            driver.delete_all_cookies()
        except Exception:
            pass
        try:
            driver.execute_script("""
                try { window.localStorage && localStorage.clear(); } catch(e) {}
                try { window.sessionStorage && sessionStorage.clear(); } catch(e) {}
            """)
        except Exception:
            pass
        try:
            driver.get("about:blank")
        except Exception:
            pass

    @staticmethod
    def _dismiss_alert_if_present(driver: webdriver.Firefox) -> None:
        try:
            a = driver.switch_to.alert
            try:
                a.dismiss()
            except Exception:
                try:
                    a.accept()
                except Exception:
                    pass
        except Exception:
            pass

    @staticmethod
    def _stop_loading(driver: webdriver.Firefox) -> None:
        try:
            driver.execute_script("window.stop();")
        except Exception:
            pass

    @staticmethod
    def _cap_fullpage_height(driver: webdriver.Firefox) -> Optional[Tuple[int, int]]:
        try:
            w = int(driver.execute_script("return Math.max(document.documentElement.clientWidth, window.innerWidth || 0) || 1920;"))
            h = int(driver.execute_script("return Math.max(document.body.scrollHeight, document.documentElement.scrollHeight, 0) || 1080;"))
            if h > _MAX_FULLPAGE_HEIGHT:
                h = _MAX_FULLPAGE_HEIGHT
            return w, h
        except Exception:
            return None

    @staticmethod
    def _try_bypass_firefox_interstitial(driver: webdriver.Firefox) -> None:
        try:
            cur = (driver.current_url or "").lower()
            if not (cur.startswith("about:certerror") or cur.startswith("about:neterror")):
                return  # Not an interstitial — skip entirely, no sleep
            driver.execute_script("""
                try {
                    const adv = document.getElementById('advancedButton') || document.querySelector('#advancedButton');
                    if (adv) adv.click();
                    const cont = document.getElementById('exceptionDialogButton') || document.querySelector('#exceptionDialogButton');
                    if (cont) cont.click();
                } catch (e) {}
            """)
            time.sleep(0.25)
        except Exception:
            pass

    @classmethod
    def capture_screenshot(
        cls,
        driver: webdriver.Firefox,
        url: str,
        follow_redirects: bool = False,
        timeout: int = DEFAULT_TIMEOUT,
        full_page: bool = False,
    ) -> Tuple[bool, Optional[bytes], str, str]:
        """Capture a screenshot.

        Returns (success, png_bytes, final_url, page_title).
        """
        png_data: Optional[bytes] = None
        try:
            try:
                driver.get(url)
            except TimeoutException:
                cls._stop_loading(driver)

            time.sleep(0.35)
            cls._try_bypass_firefox_interstitial(driver)

            js_wait_time = min(max(timeout // 4, 3), 15)
            try:
                WebDriverWait(driver, js_wait_time).until(
                    lambda d: d.execute_script("return document.readyState") in ("interactive", "complete")
                )
            except TimeoutException:
                pass

            try:
                WebDriverWait(driver, min(js_wait_time, 6)).until(
                    lambda d: cls._dom_has_meaningful_content(d)
                )
            except TimeoutException:
                pass

            page_title = ""
            try:
                page_title = (driver.title or "").strip()
            except Exception:
                pass

            final_url = url
            try:
                if follow_redirects:
                    final_url = driver.current_url or url
            except Exception:
                pass

            if full_page:
                dims = cls._cap_fullpage_height(driver)
                if dims:
                    w, h = dims
                    old = None
                    try:
                        old_w = int(driver.execute_script("return window.innerWidth;"))
                        old_h = int(driver.execute_script("return window.innerHeight;"))
                        old = (old_w, old_h)
                    except Exception:
                        pass
                    try:
                        driver.set_window_size(w, h)
                    except Exception:
                        pass
                    if hasattr(driver, "get_full_page_screenshot_as_png"):
                        try:
                            png_data = driver.get_full_page_screenshot_as_png()
                        except Exception:
                            png_data = None
                    if old:
                        try:
                            driver.set_window_size(old[0], old[1])
                        except Exception:
                            pass

            if png_data is None:
                try:
                    png_data = driver.get_screenshot_as_png()
                except UnexpectedAlertPresentException:
                    cls._dismiss_alert_if_present(driver)
                    png_data = driver.get_screenshot_as_png()

            if not png_data:
                return False, None, url, ""

            return True, png_data, final_url, page_title

        except (WebDriverException, NoSuchWindowException) as e:
            logger.debug(f"WebDriver error for {url}: {e}")
            return False, None, url, ""
        except Exception as e:
            logger.debug(f"Failed to capture screenshot for {url}: {e}")
            return False, None, url, ""
        finally:
            try:
                cls._maybe_reset_state(driver)
            except Exception:
                pass


# ═══════════════════════════════════════════════════════════════
#  Streaming capture engine
# ═══════════════════════════════════════════════════════════════
#
#  This is a GENERATOR.  It yields results in-order while printing
#  console discovery lines IMMEDIATELY as each capture completes
#  (even if out-of-order).  This gives real-time feedback during
#  the capture phase instead of the old "burst-at-the-end" pattern.
#

def _build_url(host: str, port: int, scheme: str) -> str:
    host_for_url = f"[{host}]" if is_ipv6(host) else host
    is_standard_port = (port == 80 and scheme == 'http') or (port == 443 and scheme == 'https')
    return f"{scheme}://{host_for_url}" if is_standard_port else f"{scheme}://{host_for_url}:{port}"


# Result dict keys for clarity
_R_URL     = "url"
_R_OK      = "ok"
_R_PNG     = "png"
_R_FINAL   = "final_url"
_R_TITLE   = "title"
_R_TAGS    = "tags"
_R_SCHEME  = "scheme"
_R_HOST    = "host"
_R_PORT    = "port"


def _stream_captures(
    host_port_list: List[Tuple[str, int, str]],
    progress_bar: tqdm,
    follow_redirects: bool,
    timeout: int,
    full_page: bool,
) -> Iterator[Dict]:
    """Generator: yields one dict per URL **in order**, while printing
    console discovery lines immediately as each capture finishes."""

    total = len(host_port_list)
    if total == 0:
        return

    cpu = os.cpu_count() or 2
    # Cap workers: at most _MAX_SHOT_WORKERS, at most half CPUs, at most #targets
    workers = min(_MAX_SHOT_WORKERS, max(1, cpu // 2), total)

    # Pre-warm geckodriver + options on main thread to avoid all workers
    # racing to resolve the path simultaneously (the pre-capture stall).
    WebDriverManager.warmup()

    def _task(idx: int, host: str, port: int, scheme: str, url: str) -> Dict:
        d = WebDriverManager.create_driver(timeout)
        ok, png, final_url, title = WebDriverManager.capture_screenshot(
            d, url, follow_redirects, timeout, full_page
        )
        tags = _detect_tags(title, url=url, final_url=final_url) if (ok and png) else []
        return {
            "idx": idx, _R_HOST: host, _R_PORT: port, _R_SCHEME: scheme,
            _R_URL: url, _R_OK: bool(ok and png), _R_PNG: png,
            _R_FINAL: final_url, _R_TITLE: title, _R_TAGS: tags,
        }

    it = iter(enumerate(host_port_list))
    in_flight: Dict[object, int] = {}  # future -> idx
    buf: Dict[int, Dict] = {}          # idx -> result
    next_yield = 0

    with ThreadPoolExecutor(max_workers=workers) as ex:

        def _submit() -> bool:
            try:
                i, (host, port, scheme) = next(it)
                url = _build_url(host, port, scheme)
                fut = ex.submit(_task, i, host, port, scheme, url)
                in_flight[fut] = i
                return True
            except StopIteration:
                return False

        # Seed the pool
        for _ in range(workers):
            if not _submit():
                break

        while in_flight:
            done, _ = wait(in_flight.keys(), return_when=FIRST_COMPLETED)

            for fut in done:
                idx_fallback = in_flight.pop(fut)
                progress_bar.update(1)

                try:
                    res = fut.result()
                except Exception:
                    res = {
                        "idx": idx_fallback,
                        _R_URL: "(unknown)", _R_OK: False, _R_PNG: None,
                        _R_FINAL: "", _R_TITLE: "", _R_TAGS: [],
                        _R_SCHEME: "", _R_HOST: "", _R_PORT: 0,
                    }

                # ── Live console line (IMMEDIATELY, even out-of-order) ──
                url_display = res[_R_URL]
                counter = _c_dim(f"[{res['idx']+1}/{total}]")

                if res[_R_OK]:
                    title_str = f'"{_truncate(res[_R_TITLE])}"' if res[_R_TITLE] else _c_dim("(no title)")
                    tag_str = "  " + " ".join(f"[{t}]" for t in res[_R_TAGS]) if res[_R_TAGS] else ""
                    redir = f"  -> {res[_R_FINAL]}" if res[_R_FINAL] != url_display else ""
                    _twrite(f" {_c_green('>>')} {counter} {url_display:<40} {title_str}{tag_str}{redir}")
                else:
                    _twrite(f" {_c_red('!!')} {counter} {url_display:<40} {_c_red('capture failed')}")

                buf[res["idx"]] = res

                # Submit replacement work right away to keep the pool full
                _submit()

            # Yield results in order (file output must be ordered)
            while next_yield in buf:
                yield buf.pop(next_yield)
                next_yield += 1


# ═══════════════════════════════════════════════════════════════
#  Document generator
# ═══════════════════════════════════════════════════════════════

_HTML_HEADER = """<!DOCTYPE html><html lang="en"><head><title>NetGazer Scan Report</title><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1"><style>
* {{ box-sizing: border-box; }} body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f3f4f6; margin: 0; padding: 24px; color: #1f2937; }}
.header {{ text-align: center; margin: 0 auto 14px; max-width: 1400px; padding: 24px; background: #fff; border-radius: 10px; box-shadow: 0 1px 3px rgba(0,0,0,.08); }} .header h1 {{ margin: 0 0 10px; font-size: 22px; font-weight: 700; }}
.meta {{ color: #6b7280; font-size: 13px; line-height: 1.7; }} .meta b {{ color: #374151; font-weight: 600; }}
.toolbar {{ margin: 0 auto 10px; max-width: 1400px; display: grid; gap: 10px; }} .toolbar input {{ width: 100%; padding: 10px 14px; border: 1px solid #d1d5db; border-radius: 8px; font-size: 14px; outline: none; background: #fff; transition: border .15s; }}
.toolbar input:focus {{ border-color: #3b82f6; box-shadow: 0 0 0 3px rgba(59,130,246,.15); }}
.tagbar {{ display: flex; flex-wrap: wrap; gap: 8px; align-items: center; }}
.tagbtn {{ border: 1px solid #d1d5db; background: #fff; color: #374151; padding: 6px 10px; border-radius: 999px; font-size: 12px; cursor: pointer; user-select: none; }}
.tagbtn.active {{ border-color: #111827; color: #111827; box-shadow: 0 0 0 3px rgba(17,24,39,.08); }}
.tagbtn .dot {{ display: inline-block; width: 8px; height: 8px; border-radius: 999px; margin-right: 6px; vertical-align: middle; }} .status {{ color: #6b7280; font-size: 12px; }}
table {{ margin: 0 auto; border-collapse: collapse; width: 100%; max-width: 1400px; background: #fff; border-radius: 10px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,.08); }}
thead {{ background: #f9fafb; }} th {{ font-weight: 600; text-align: left; font-size: 12px; text-transform: uppercase; letter-spacing: .5px; color: #6b7280; padding: 12px 16px; border-bottom: 2px solid #e5e7eb; }}
td {{ padding: 14px 16px; vertical-align: top; border-bottom: 1px solid #f3f4f6; }} td:first-child {{ width: 320px; min-width: 240px; }}
.url-cell {{ font-size: 13px; line-height: 1.5; }} .url-link {{ color: #2563eb; text-decoration: none; word-break: break-all; font-weight: 500; }} .url-link:hover {{ text-decoration: underline; }}
.scheme-http {{ display: inline-block; font-size: 9px; font-weight: 700; color: #dc2626; background: #fef2f2; padding: 1px 5px; border-radius: 3px; margin-right: 4px; vertical-align: middle; }}
.scheme-https {{ display: inline-block; font-size: 9px; font-weight: 700; color: #16a34a; background: #f0fdf4; padding: 1px 5px; border-radius: 3px; margin-right: 4px; vertical-align: middle; }}
.page-title {{ color: #6b7280; font-size: 12px; margin-top: 5px; }} .page-title em {{ font-style: italic; }}
.redirect {{ color: #9ca3af; font-size: 11px; margin-top: 3px; word-break: break-all; }}
.tag {{ display: inline-block; padding: 2px 7px; border-radius: 4px; font-size: 10px; font-weight: 700; margin-top: 6px; margin-right: 3px; color: #fff; text-transform: uppercase; letter-spacing: .3px; }}
img {{ max-width: 100%; height: auto; border-radius: 6px; cursor: zoom-in; display: block; transition: opacity .2s; }}
.footer {{ text-align: center; margin: 20px auto 0; max-width: 1400px; padding: 16px; color: #9ca3af; font-size: 12px; }}
details {{ max-width: 1400px; margin: 12px auto 0; background: #fff; border-radius: 10px; box-shadow: 0 1px 3px rgba(0,0,0,.08); padding: 10px 14px; }}
summary {{ cursor: pointer; color: #374151; font-weight: 600; }} pre {{ white-space: pre-wrap; word-break: break-all; color: #6b7280; font-size: 12px; margin: 10px 0 0; }}
#overlay {{ display: none; position: fixed; inset: 0; background: rgba(0,0,0,.92); z-index: 1000; cursor: zoom-out; justify-content: center; align-items: center; }}
#overlay img {{ max-width: 95%; max-height: 95%; object-fit: contain; }}
@media print {{ .toolbar, #overlay, details {{ display: none !important; }} body {{ background: #fff; padding: 0; }} table {{ box-shadow: none; }} img {{ cursor: default; }} }}
</style></head><body>
<div class="header"><h1>NetGazer Scan Report</h1><div class="meta">{meta_html}</div></div>
<div class="toolbar"><input type="text" id="filterInput" onkeyup="filterRows()" placeholder="Filter by URL, title, or tag..."><div class="tagbar" id="tagBar">{tagbar_html}<span class="status" id="statusLine"></span></div></div>
<table><thead><tr><th>Web Server</th><th>Screenshot</th></tr></thead><tbody id="tbody">
"""

_HTML_FOOTER = """</tbody></table>
{failed_block}
<div class="footer">{footer_html}</div>
<div id="overlay" onclick="closeZoom()"><img id="zoomedImg" alt="Zoomed screenshot"></div>
<script>
let activeTag = "";
function zoom(img){{ document.getElementById('zoomedImg').src = img.src; document.getElementById('overlay').style.display = 'flex'; }}
function closeZoom(){{ document.getElementById('overlay').style.display = 'none'; }}
document.addEventListener('keydown', function(e){{ if (e.key === 'Escape') closeZoom(); }});
function setTag(tag){{ if (activeTag === tag){{ activeTag = ""; }} else {{ activeTag = tag; }}
document.querySelectorAll('.tagbtn').forEach(btn => {{ btn.classList.toggle('active', btn.dataset.tag === activeTag); }}); filterRows(); }}
function filterRows(){{ const term = (document.getElementById('filterInput').value || "").toLowerCase();
const rows = document.querySelectorAll('#tbody tr'); let shown = 0; rows.forEach(row => {{
const hay = (row.dataset.search || ""); const tags = (row.dataset.tags || "");
const okTerm = !term || hay.includes(term); const okTag = !activeTag || tags.split(',').includes(activeTag);
const show = okTerm && okTag; row.style.display = show ? '' : 'none'; if (show) shown++; }});
document.getElementById('statusLine').textContent = `Showing ${{shown}} / ${{rows.length}}`; }}
filterRows();
</script></body></html>"""


class DocumentGenerator:

    @staticmethod
    def build_url(host: str, port: int, scheme: str) -> str:
        return _build_url(host, port, scheme)

    # ───────────────────────── HTML ─────────────────────────

    @classmethod
    def generate_html(
        cls,
        output_file: str,
        hosts_to_capture: List[Tuple[str, List[Tuple[int, str]]]],
        progress_bar: tqdm,
        follow_redirects: bool = False,
        timeout: int = DEFAULT_TIMEOUT,
        full_page: bool = False,
        scan_meta: Optional[Dict] = None
    ) -> Tuple[int, int, Counter, List[str]]:

        meta = scan_meta or {}
        host_port_list: List[Tuple[str, int, str]] = [
            (host, port, scheme)
            for host, port_scheme_list in hosts_to_capture
            for port, scheme in port_scheme_list
        ]

        failures: List[str] = []
        captured_n = 0
        tag_counter: Counter = Counter()

        # Partial metadata (counts filled in later)
        meta_parts_pre = []
        if meta.get("timestamp"):
            meta_parts_pre.append(f"<b>Date:</b> {html.escape(str(meta['timestamp']))}")
        if meta.get("total_hosts"):
            meta_parts_pre.append(f"<b>Targets:</b> {int(meta['total_hosts'])} hosts")
        if meta.get("ports"):
            meta_parts_pre.append(f"<b>Ports:</b> {html.escape(', '.join(map(str, meta['ports'])))}")

        # Stream rows to a temp file as they arrive (memory-efficient)
        rows_tmp = tempfile.NamedTemporaryFile("w", encoding="utf-8", delete=False, suffix=".rows.html")
        rows_tmp_path = Path(rows_tmp.name)

        try:
            # ── Stream captures, writing each row immediately ──
            for r in _stream_captures(host_port_list, progress_bar, follow_redirects, timeout, full_page):
                if r[_R_OK] and r[_R_PNG]:
                    tags = r[_R_TAGS]
                    for t in tags:
                        tag_counter[t] += 1

                    b64_data = ""
                    try:
                        b64_data = base64.b64encode(r[_R_PNG]).decode('utf-8')
                    except Exception:
                        b64_data = ""

                    url = r[_R_URL]
                    final_url = r[_R_FINAL]
                    title = r[_R_TITLE]
                    scheme = r[_R_SCHEME]

                    safe_url_attr = html.escape(url, quote=True)
                    safe_url_text = html.escape(url)
                    safe_final = html.escape(final_url, quote=True)
                    search_blob = (f"{url} {title} {' '.join(tags)}").lower()
                    safe_search = html.escape(search_blob, quote=True)
                    safe_tags_attr = html.escape(",".join(tags), quote=True)

                    scheme_badge = ('<span class="scheme-https">HTTPS</span>'
                                    if scheme == "https"
                                    else '<span class="scheme-http">HTTP</span>')
                    link = f'<a href="{safe_url_attr}" target="_blank" rel="noopener" class="url-link">{safe_url_text}</a>'

                    cell_parts = [f'{scheme_badge}{link}']
                    if title:
                        cell_parts.append(f'<div class="page-title"><em>{html.escape(_truncate(title))}</em></div>')
                    if final_url != url:
                        cell_parts.append(f'<div class="redirect">→ {safe_final}</div>')
                    if tags:
                        badge_html = " ".join(
                            f'<span class="tag" style="background:{_TAG_BADGE_COLORS.get(t, "#6b7280")}">{html.escape(t)}</span>'
                            for t in tags
                        )
                        cell_parts.append(badge_html)

                    url_cell = f'<td class="url-cell">{"".join(cell_parts)}</td>'
                    img_cell = (
                        f'<td><img src="data:image/png;base64,{b64_data}" alt="Screenshot of {safe_url_text}" '
                        f'loading="lazy" decoding="async" onclick="zoom(this)"></td>'
                    )
                    rows_tmp.write(f'    <tr data-search="{safe_search}" data-tags="{safe_tags_attr}">{url_cell}{img_cell}</tr>\n')
                    captured_n += 1
                else:
                    failures.append(r[_R_URL])

            rows_tmp.close()

        except Exception:
            try:
                rows_tmp.close()
            except Exception:
                pass
            raise

        # ── Assemble final HTML ──
        total_n = len(host_port_list)
        failed_n = len(failures)

        meta_parts2 = list(meta_parts_pre)
        meta_parts2.append(f"<b>Captured:</b> {captured_n}/{total_n} screenshots")
        if failed_n:
            meta_parts2.append(f"<b style='color:#dc2626'>Failed:</b> {failed_n}")
        if tag_counter:
            tag_summary = ", ".join(f"{cnt}× {tag}" for tag, cnt in tag_counter.most_common())
            meta_parts2.append(f"<b>Flagged:</b> {html.escape(tag_summary)}")
        meta_html = " &nbsp;|&nbsp; ".join(meta_parts2)

        tagbar_parts = []
        if tag_counter:
            tagbar_parts.append('<span class="tagbtn active" data-tag="" onclick="setTag(\'\')">All</span>')
            for tag, cnt in tag_counter.most_common():
                color = _TAG_BADGE_COLORS.get(tag, "#6b7280")
                tagbar_parts.append(
                    f'<span class="tagbtn" data-tag="{html.escape(tag)}" onclick="setTag(\'{html.escape(tag)}\')">'
                    f'<span class="dot" style="background:{color}"></span>{html.escape(tag)} ({cnt})</span>'
                )
        tagbar_html = "\n".join(tagbar_parts)

        footer_parts = [f"Generated by NetGazer — {captured_n} screenshots captured"]
        if failed_n:
            footer_parts.append(f"{failed_n} failed")
        footer_html = " | ".join(footer_parts)

        failed_block = ""
        if failures:
            failed_block = (
                f"<details><summary>Failed URLs ({len(failures)})</summary>"
                f"<pre>{html.escape(chr(10).join(failures))}</pre></details>"
            )

        out_path = Path(output_file)
        tmp_out = out_path.with_suffix(out_path.suffix + ".tmp")
        try:
            with tmp_out.open("w", encoding="utf-8") as out:
                out.write(_HTML_HEADER.format(meta_html=meta_html, tagbar_html=tagbar_html))
                with rows_tmp_path.open("r", encoding="utf-8") as rf:
                    shutil.copyfileobj(rf, out)
                out.write(_HTML_FOOTER.format(failed_block=failed_block, footer_html=footer_html))
            _atomic_replace(tmp_out, out_path)
        except IOError as e:
            logger.error(f"Failed to write HTML file {output_file}: {e}")
            sys.exit(1)
        finally:
            for p in (rows_tmp_path, tmp_out):
                try:
                    p.unlink(missing_ok=True)
                except Exception:
                    pass

        return captured_n, failed_n, tag_counter, failures

    # ───────────────────────── DOCX ─────────────────────────

    @staticmethod
    def generate_docx(
        output_file: str,
        hosts_to_capture: List[Tuple[str, List[Tuple[int, str]]]],
        progress_bar: tqdm,
        follow_redirects: bool = False,
        timeout: int = DEFAULT_TIMEOUT,
        full_page: bool = False,
        scan_meta: Optional[Dict] = None
    ) -> Tuple[int, int, Counter, List[str]]:

        captured = 0
        failed = 0
        tag_counter: Counter = Counter()
        failure_urls: List[str] = []

        host_port_list: List[Tuple[str, int, str]] = [
            (host, port, scheme)
            for host, port_scheme_list in hosts_to_capture
            for port, scheme in port_scheme_list
        ]

        meta = scan_meta or {}
        out_path = Path(output_file)
        tmp_out = out_path.with_suffix(out_path.suffix + ".tmp")

        try:
            doc = docx.Document()  # type: ignore[union-attr]

            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_paragraph.add_run("Web Server Screenshots")
            title_run.bold = True
            title_run.font.size = Pt(24)

            meta_parts = []
            if meta.get("timestamp"):
                meta_parts.append(f"Date: {meta['timestamp']}")
            if meta.get("total_hosts"):
                meta_parts.append(f"Targets: {meta['total_hosts']} hosts")
            if meta.get("ports"):
                meta_parts.append(f"Ports: {', '.join(map(str, meta['ports']))}")
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                meta_run = meta_para.add_run(" | ".join(meta_parts))
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            doc.add_paragraph("")

            table = doc.add_table(rows=1, cols=2)
            try:
                table.style = 'Light Grid Accent 1'
            except Exception:
                try:
                    table.style = 'Table Grid'
                except Exception:
                    pass
            table.autofit = False
            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.0)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Web Request Info'
            hdr_cells[1].text = 'Web Screenshot'
            for cell in hdr_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # ── Stream captures, adding rows as they arrive ──
            for r in _stream_captures(host_port_list, progress_bar, follow_redirects, timeout, full_page):
                if not r[_R_OK] or not r[_R_PNG]:
                    failed += 1
                    failure_urls.append(r[_R_URL])
                    continue

                tags = r[_R_TAGS]
                for t in tags:
                    tag_counter[t] += 1

                url = r[_R_URL]
                final_url = r[_R_FINAL]
                title = r[_R_TITLE]

                row_cells = table.add_row().cells
                url_cell = row_cells[0]
                url_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                url_para = url_cell.paragraphs[0]
                url_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                url_run = url_para.add_run(url)
                url_run.font.size = Pt(9)

                if final_url != url:
                    redir_para = url_cell.add_paragraph()
                    redir_run = redir_para.add_run(f"→ {final_url}")
                    redir_run.font.size = Pt(8)
                    redir_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                if title:
                    title_para = url_cell.add_paragraph()
                    t_run = title_para.add_run(f'"{_truncate(title)}"')
                    t_run.italic = True
                    t_run.font.size = Pt(8)
                    t_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                if tags:
                    tag_para = url_cell.add_paragraph()
                    tag_run = tag_para.add_run(f'[{", ".join(tags)}]')
                    tag_run.bold = True
                    tag_run.font.size = Pt(8)
                    tag_run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)

                with BytesIO(r[_R_PNG]) as img_buf:
                    pic_run = row_cells[1].paragraphs[0].add_run()
                    pic_run.add_picture(img_buf, width=Inches(4.0))

                captured += 1

            if failure_urls:
                doc.add_page_break()
                p = doc.add_paragraph("Failed URLs")
                if p.runs:
                    p.runs[0].bold = True
                for u in failure_urls:
                    try:
                        doc.add_paragraph(u, style='List Bullet')
                    except Exception:
                        doc.add_paragraph(f"- {u}")

            section = doc.sections[0]
            footer = section.footer
            if footer.paragraphs:
                footer_paragraph = footer.paragraphs[0]
            else:
                footer_paragraph = footer.add_paragraph()
            footer_parts = ["Generated by NetGazer"]
            if meta.get("timestamp"):
                footer_parts.append(meta["timestamp"])
            footer_parts.append(f"{captured} screenshots")
            footer_paragraph.text = " | ".join(footer_parts)
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.save(str(tmp_out))
            _atomic_replace(tmp_out, out_path)

        except Exception as e:
            logger.error(f"Failed to generate DOCX file {output_file}: {e}")
            _twrite(f" {_c_red('!!')} DOCX generation error: {e}")
            try:
                if tmp_out.exists():
                    tmp_out.unlink()
            except Exception:
                pass

        return captured, failed, tag_counter, failure_urls


# ═══════════════════════════════════════════════════════════════
#  CLI parsing helpers
# ═══════════════════════════════════════════════════════════════

def comma_separated_ints(value: str) -> List[int]:
    try:
        ports = [int(p.strip()) for p in value.split(',') if p.strip()]
        if not ports:
            raise ValueError("No valid ports provided")

        invalid_ports = [p for p in ports if not 1 <= p <= 65535]
        if invalid_ports:
            raise ValueError(f"Ports out of valid range (1-65535): {invalid_ports}")

        return sorted(set(ports))
    except ValueError as e:
        raise argparse.ArgumentTypeError(f"Invalid port specification: {e}")


def check_dependencies(output_file: str) -> None:
    missing = []
    if webdriver is None:
        missing.append("selenium")
    if tqdm is None:
        missing.append("tqdm")

    out_suffix = Path(output_file).suffix.lower()
    if out_suffix == ".docx" and docx is None:
        missing.append("python-docx")

    if _resolve_geckodriver_path() is None and GeckoDriverManager is None:
        missing.append("webdriver-manager")

    if missing:
        print(f"Missing required libraries: {', '.join(missing)}")
        print(f"Install with: pip install {' '.join(missing)}")
        sys.exit(1)


def initial_checks(output_file: str) -> None:
    if hasattr(os, "geteuid") and os.geteuid() == 0:
        print(f"  {_c_yellow('[!]')} Running as root — browser may behave unexpectedly")
    check_dependencies(output_file)


# ═══════════════════════════════════════════════════════════════
#  Scan phases
# ═══════════════════════════════════════════════════════════════

def perform_port_scan(hosts: List[str], ports: List[int]) -> List[Tuple[str, List[Tuple[int, str]]]]:
    host_order = {h: i for i, h in enumerate(hosts)}
    ordered_results: List[Tuple[int, Tuple[str, List[Tuple[int, str]]]]] = []

    t0 = time.perf_counter()
    max_workers = min(MAX_PORT_WORKERS, max(1, len(hosts)))

    with tqdm(total=len(hosts), unit='host', desc="  Port scan", leave=False) as pbar:  # type: ignore[union-attr]
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            it = iter(hosts)
            in_flight: Dict[object, str] = {}

            def _submit(h: str) -> None:
                fut = executor.submit(PortScanner.scan_host, h, ports)
                in_flight[fut] = h

            for _ in range(max_workers):
                try:
                    _submit(next(it))
                except StopIteration:
                    break

            while in_flight:
                done, _ = wait(in_flight.keys(), return_when=FIRST_COMPLETED)
                for fut in done:
                    host = in_flight.pop(fut)
                    pbar.update(1)
                    try:
                        result = fut.result()
                        if result:
                            ordered_results.append((host_order.get(host, 10**9), result))
                    except Exception as e:
                        logger.error(f"Error during port scan for {host}: {e}")
                    try:
                        _submit(next(it))
                    except StopIteration:
                        pass

    ordered_results.sort(key=lambda x: x[0])
    hosts_to_capture = [r for _, r in ordered_results]

    dt = time.perf_counter() - t0
    total_urls = sum(len(pl) for _, pl in hosts_to_capture)
    hosts_found = len(hosts_to_capture)

    port_counter: Counter = Counter()
    for _, pl in hosts_to_capture:
        for p, _ in pl:
            port_counter[p] += 1

    print(f"  Found {_c_bold(str(total_urls))} web server{'s' if total_urls != 1 else ''} "
          f"on {hosts_found}/{len(hosts)} hosts  "
          f"{_c_dim(_fmt_elapsed(dt))}")

    if port_counter:
        top = ", ".join(f"{p}({c})" for p, c in port_counter.most_common(6))
        print(f"  Common: {top}")

    with _DNS_FAILS_LOCK:
        dns_fail_n = len(_DNS_FAILS)
    if dns_fail_n:
        print(f"  DNS failed: {_c_yellow(str(dns_fail_n))} host{'s' if dns_fail_n != 1 else ''}")

    return hosts_to_capture


def _write_failed_list(output_file: str, failures: List[str]) -> None:
    if not failures:
        return
    try:
        p = Path(output_file)
        out = p.with_suffix(p.suffix + ".failed.txt")
        out.write_text("\n".join(failures) + "\n", encoding="utf-8")
    except Exception:
        pass


def perform_screenshot_capture(
    hosts_to_capture: List[Tuple[str, List[Tuple[int, str]]]],
    output_file: str,
    follow_redirects: bool,
    timeout: int,
    full_page: bool,
    scan_meta: Optional[Dict] = None
) -> Tuple[int, int, Counter]:
    total_screenshots = sum(len(port_list) for _, port_list in hosts_to_capture)
    print(f"\n  Screenshot Capture ({total_screenshots} URLs, {timeout}s timeout)")

    captured, failed_count = 0, 0
    tag_counter: Counter = Counter()
    failures: List[str] = []

    t0 = time.perf_counter()
    try:
        output_path = Path(output_file)
        with tqdm(total=total_screenshots, unit='shot', desc="  Capturing", leave=False) as pbar:  # type: ignore[union-attr]
            if output_path.suffix.lower() == '.docx':
                captured, failed_count, tag_counter, failures = DocumentGenerator.generate_docx(
                    output_file, hosts_to_capture, pbar,
                    follow_redirects, timeout, full_page, scan_meta
                )
            else:
                captured, failed_count, tag_counter, failures = DocumentGenerator.generate_html(
                    output_file, hosts_to_capture, pbar,
                    follow_redirects, timeout, full_page, scan_meta
                )
    finally:
        # Cleanup in the background so the user sees the summary instantly
        WebDriverManager.cleanup_all()

    dt = time.perf_counter() - t0
    if failures:
        _write_failed_list(output_file, failures)

    fail_str = f"  ({_c_red(str(failed_count) + ' failed')})" if failed_count else ""
    print(f"  Captured {_c_bold(str(captured))}/{total_screenshots} screenshots{fail_str}  {_c_dim(_fmt_elapsed(dt))}")

    return captured, failed_count, tag_counter


# ═══════════════════════════════════════════════════════════════
#  Argument validation
# ═══════════════════════════════════════════════════════════════

def validate_args(args: argparse.Namespace) -> None:
    output_path = Path(args.output_file)
    if output_path.suffix.lower() not in ['.html', '.docx']:
        print("Error: Output file must have .html or .docx extension")
        sys.exit(1)

    if args.timeout <= 0:
        print("Error: Timeout must be positive")
        sys.exit(1)

    output_dir = output_path.parent
    if not output_dir.exists():
        try:
            output_dir.mkdir(parents=True, exist_ok=True)
        except (OSError, PermissionError):
            print(f"Error: Cannot create output directory {output_dir}")
            sys.exit(1)

    if not os.access(output_dir, os.W_OK):
        print(f"Error: No write permission for output directory {output_dir}")
        sys.exit(1)


# ═══════════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════════

def main() -> None:
    parser = argparse.ArgumentParser(
        description='Scan and capture web server screenshots.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s hosts.txt results.html
  %(prog)s 192.168.1.0/24 scan.docx --timeout 45 --redirect
  %(prog)s targets.txt output.html --ports 8080,8443,9000
  %(prog)s example.com results.html --ports 80,443,8080,8443,9090
        """
    )

    parser.add_argument('input', help='Input hosts file, CIDR notation, or IP range / URL')
    parser.add_argument('output_file', help='Output file name (.docx or .html)')
    parser.add_argument('--timeout', '-t', type=int, default=DEFAULT_TIMEOUT,
                        help=f'Page load timeout in seconds (default: {DEFAULT_TIMEOUT})')
    parser.add_argument('--redirect', '-r', action='store_true',
                        help='Label final destination if page redirects (browser still follows redirects).')
    parser.add_argument('--ports', '-p', type=comma_separated_ints, default=DEFAULT_PORTS,
                        help='Comma-separated list of ports to scan (default: 80,443)')
    parser.add_argument('--full-page', action='store_true',
                        help='Capture full-page screenshots (Firefox-only feature).')

    args = parser.parse_args()

    validate_args(args)
    initial_checks(args.output_file)

    hosts = NetworkScanner.parse_hosts(args.input)
    if not hosts:
        print("No valid hosts found to scan.")
        sys.exit(1)

    expanded_hosts = _dedupe_preserve_order(hosts)
    if not expanded_hosts:
        print("No hosts to scan after expansion.")
        sys.exit(1)

    port_str = ", ".join(map(str, args.ports))
    redir_label = "label" if args.redirect else "off"
    fp_label = "  Full-page: on" if args.full_page else ""

    print(f"\n{'─' * 64}")
    print(f"  NetGazer — Web Server Screenshot Scanner")
    print(f"  Targets  : {len(expanded_hosts)} host{'s' if len(expanded_hosts) != 1 else ''}")
    print(f"  Ports    : {port_str}")
    print(f"  Output   : {args.output_file}")
    print(f"  Timeout  : {args.timeout}s   Redirects: {redir_label}{fp_label}")
    print(f"{'─' * 64}")

    t_total = time.perf_counter()

    print(f"\n  Port Scan")
    hosts_to_capture = perform_port_scan(expanded_hosts, args.ports)

    if not hosts_to_capture:
        total_dt = time.perf_counter() - t_total
        print(f"\n  No web servers found on scanned hosts.")
        print(f"  Tip: try --ports 8080,8443,9090 if non-standard ports are in use")
        print(f"  Total: {_fmt_elapsed(total_dt)}")
        print(f"{'─' * 64}")
        return

    scan_meta = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_hosts": len(expanded_hosts),
        "hosts_with_ports": len(hosts_to_capture),
        "ports": args.ports,
        "timeout": args.timeout,
        "follow_redirects": args.redirect,
    }

    captured, failed_count, tag_counter = perform_screenshot_capture(
        hosts_to_capture, args.output_file, args.redirect,
        args.timeout, args.full_page, scan_meta
    )

    total_dt = time.perf_counter() - t_total

    print(f"\n{'─' * 64}")
    print(f"  Saved: {_c_bold(args.output_file)} ({captured} screenshot{'s' if captured != 1 else ''})")

    if tag_counter:
        tag_parts = [f"{cnt}× {tag}" for tag, cnt in tag_counter.most_common()]
        print(f"  Flagged: {', '.join(tag_parts)}")

    if failed_count:
        print(f"  Failed: {failed_count} capture{'s' if failed_count != 1 else ''}")
        print(f"  Retry list: {Path(args.output_file).with_suffix(Path(args.output_file).suffix + '.failed.txt')}")

    print(f"  Total: {_fmt_elapsed(total_dt)}")
    print(f"{'─' * 64}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python netgazer.py <hosts_file | IP/CIDR/range/domain/URL> <output_file> [--timeout SECONDS] [--redirect] [--full-page] [--ports PORT1,PORT2,...]")
        print("Examples:")
        print("  python netgazer.py hosts.txt results.html")
        print("  python netgazer.py 192.168.1.0/24 scan.docx --timeout 45 --redirect --full-page")
        print("  python netgazer.py targets.txt output.html --ports 8080,8443,9000")
        print("  python netgazer.py example.com results.html --ports 80,443,8080,8443,9090")
    else:
        main()
